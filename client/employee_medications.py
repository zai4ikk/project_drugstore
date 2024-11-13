from PyQt6.QtWidgets import QWidget, QVBoxLayout, QLabel, QPushButton, QScrollArea, QDialog, QLineEdit
from PyQt6.QtCore import Qt
import mysql.connector
import os
from openpyxl import Workbook
from docx import Document
from provider_list import ProviderList  # Подключение вашего модуля ProviderList


class MedicationList(QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Список медикаментов")

        self.layout = QVBoxLayout()
        self.add_medication_button = QPushButton("Добавить медикамент")
        self.add_medication_button.clicked.connect(self.add_medication_dialog)
        self.layout.addWidget(self.add_medication_button)

        # Кнопки для отчетов и просмотра заказов
        provider_button = QPushButton("Поставщики медикаментов")
        med_button = QPushButton("Отчет о медикаментах")
        ord_button = QPushButton("Отчет о заказах")
        view_orders_button = QPushButton("Посмотреть заказы")

        self.layout.addWidget(provider_button)
        self.layout.addWidget(med_button)
        self.layout.addWidget(ord_button)
        self.layout.addWidget(view_orders_button)

        # Обработчики событий для кнопок
        provider_button.clicked.connect(self.show_provider_list)
        med_button.clicked.connect(self.generate_word_report)
        ord_button.clicked.connect(self.generate_excel_report)
        view_orders_button.clicked.connect(self.view_orders)

        # Создание области прокрутки
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        content_widget = QWidget()
        self.scroll_layout = QVBoxLayout(content_widget)
        self.scroll_area.setWidget(content_widget)
        self.layout.addWidget(self.scroll_area)

        self.setLayout(self.layout)

        # Подключение к базе данных MySQL
        self.db = mysql.connector.connect(
            host="localhost",
            user="root",
            password="sas",
            database="pharmacy"
        )

    def view_orders(self):
        # Очистка существующего содержимого области прокрутки
        for i in reversed(range(self.scroll_layout.count())):
            widget = self.scroll_layout.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        # Запрос на получение данных о заказах
        cursor = self.db.cursor()
        cursor.execute("SELECT ordered.id, Medicament.Name, ordered.Pricerealization, ordered.Quantity "
                       "FROM ordered "
                       "INNER JOIN medicament ON ordered.IDMedicament = medicament.id")
        orders = cursor.fetchall()

        # Отображение заказов
        for order in orders:
            order_info = f"""
                Код заказа: {order[0]}
                Наименование медикамента: {order[1]}
                Цена реализации: {order[2]}
                Количество: {order[3]}
            """

            order_label = QLabel(order_info)
            order_label.setWordWrap(True)
            self.scroll_layout.addWidget(order_label)

        # Фиксация изменений и закрытие курсора
        self.db.commit()
        cursor.close()

    def generate_word_report(self):
        # Создание нового документа Word
        doc = Document()
        doc.add_heading('Отчет о медикаментах', level=1)

        # Получение данных о медикаментах
        cursor = self.db.cursor()
        cursor.execute("SELECT Medicament.Name, Medicament.Pricepurchase, "
                       "Medicament.Pricerealization, Assignationmedicament.Groupmedicament, "
                       "Assignationmedicament.Description, Diseases.Illness "
                       "FROM Medicament "
                       "INNER JOIN Assignationmedicament ON Medicament.IDAssignationmedicament = Assignationmedicament.id "
                       "LEFT JOIN Diseases ON Medicament.IDAssignationmedicament = Diseases.IDAssignationmedicament")
        medications = cursor.fetchall()

        # Добавление информации о каждом медикаменте в документ Word
        for medication in medications:
            doc.add_heading(f"{medication[0]}", level=2)
            doc.add_paragraph(f"Цена закупки: {medication[1]}")
            doc.add_paragraph(f"Цена реализации: {medication[2]}")
            doc.add_paragraph(f"Группа медикаментов: {medication[3]}")
            doc.add_paragraph(f"Описание: {medication[4]}")
            doc.add_paragraph(f"Заболевание: {medication[5]}")
            doc.add_paragraph("\n")

        # Сохранение документа Word
        doc.save('отчет_о_медикаментах.docx')

        # Открытие файла Word с использованием приложения по умолчанию
        os.system('start отчет_о_медикаментах.docx')

    def generate_excel_report(self):
        # Создание новой книги Excel и выбор активного листа
        workbook = Workbook()
        sheet = workbook.active

        # Добавление заголовков в лист Excel
        headers = ["Код заказа", "Наименование медикамента", "Цена реализации", "Количество"]
        sheet.append(headers)

        # Извлечение данных о заказах из базы данных
        cursor = self.db.cursor()
        cursor.execute("SELECT ordered.id, medicament.Name, ordered.Pricerealization, ordered.Quantity "
                       "FROM ordered "
                       "INNER JOIN medicament ON ordered.IDMedicament = medicament.id")
        orders_data = cursor.fetchall()

        # Добавление данных в лист Excel
        for order in orders_data:
            sheet.append(order)

        # Сохранение книги Excel
        workbook.save('отчет_о_заказах.xlsx')

        # Открытие файла Excel с использованием приложения по умолчанию
        os.system('start отчет_о_заказах.xlsx')

    def add_medication_dialog(self):
        add_medication_dialog = AddMedicationDialog(self)
        add_medication_dialog.exec()

    def add_medication(self, name, price_purchase, price_realization, group, description, illness, image_path):
        # Добавление медикамента в базу данных
        cursor = self.db.cursor()
        cursor.execute(
            "INSERT INTO Medicament (Name, IDAssignationmedicament, Unitsofmeasurement, Pricepurchase, Pricerealization) "
            "VALUES (%s, (SELECT id FROM Assignationmedicament WHERE Groupmedicament = %s), %s, %s, %s)",
            (name, group, 'шт', price_purchase, price_realization)
        )
        medication_id = cursor.lastrowid

        # Добавление информации о болезни в таблицу Diseases
        cursor.execute(
            "INSERT INTO Diseases (Illness, IDAssignationmedicament) VALUES (%s, (SELECT id FROM Assignationmedicament WHERE Groupmedicament = %s))",
            (illness, group)
        )

        # Фиксация изменений в базе данных
        self.db.commit()

        # Отображение информации о добавленном медикаменте на форме (не реализовано в полном объеме)

    def show_provider_list(self):
        if not hasattr(self, 'provider_list'):
            self.provider_list = ProviderList()
        self.provider_list.show()


class AddMedicationDialog(QDialog):
    def __init__(self, medication_list):
        super().__init__()
        self.setWindowTitle('Добавить медикамент')

        self.medication_list = medication_list

        self.name_label = QLabel('Наименование:')
        self.name_input = QLineEdit()

        self.price_purchase_label = QLabel('Цена закупки:')
        self.price_purchase_input = QLineEdit()

        self.price_realization_label = QLabel('Цена реализации:')
        self.price_realization_input = QLineEdit()

        self.group_label = QLabel('Группа медикаментов:')
        self.group_input = QLineEdit()

        self.description_label = QLabel('Описание:')
        self.description_input = QLineEdit()

        self.illness_label = QLabel('Заболевание:')
        self.illness_input = QLineEdit()

        self.image_path_label = QLabel('Путь к изображению:')
        self.image_path_input = QLineEdit()

        self.add_medication_button = QPushButton('Добавить медикамент')
        self.add_medication_button.clicked.connect(self.add_medication)

        layout = QVBoxLayout(self)
        layout.addWidget(self.name_label)
        layout.addWidget(self.name_input)
        layout.addWidget(self.price_purchase_label)
        layout.addWidget(self.price_purchase_input)
        layout.addWidget(self.price_realization_label)
        layout.addWidget(self.price_realization_input)
        layout.addWidget(self.group_label)
        layout.addWidget(self.group_input)
        layout.addWidget(self.description_label)
        layout.addWidget(self.description_input)
        layout.addWidget(self.illness_label)
        layout.addWidget(self.illness_input)
        layout.addWidget(self.image_path_label)
        layout.addWidget(self.image_path_input)
        layout.addWidget(self.add_medication_button)

    def add_medication(self):
        name = self.name_input.text()
        price_purchase = self.price_purchase_input.text()
        price_realization = self.price_realization_input.text()
        group = self.group_input.text()
        description = self.description_input.text()
        illness = self.illness_input.text()
        image_path = self.image_path_input.text()

        self.medication_list.add_medication(name, price_purchase, price_realization, group, description, illness, image_path)
        self.close()
